"""
export.py
---------
Renders a resume dict into downloadable .docx / .pdf files in memory.
"""

import io
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet


def build_docx(resume: dict) -> io.BytesIO:
    doc = Document()
    doc.add_heading(resume["full_name"], level=0)

    contact_line = " | ".join(filter(None, [resume.get("email"), resume.get("phone")]))
    if contact_line:
        doc.add_paragraph(contact_line)

    if resume.get("summary"):
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(resume["summary"])

    if resume.get("skills"):
        doc.add_heading("Skills", level=1)
        doc.add_paragraph(", ".join(resume["skills"]))

    if resume.get("experience"):
        doc.add_heading("Experience", level=1)
        for job in resume["experience"]:
            p = doc.add_paragraph()
            run = p.add_run(f"{job.get('title', '')} — {job.get('company', '')}")
            run.bold = True
            if job.get("duration"):
                doc.add_paragraph(job["duration"])
            for bullet in job.get("bullets", []):
                doc.add_paragraph(bullet, style="List Bullet")

    if resume.get("education"):
        doc.add_heading("Education", level=1)
        for edu in resume["education"]:
            doc.add_paragraph(edu)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def build_pdf(resume: dict) -> io.BytesIO:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    elements = []

    elements.append(Paragraph(resume["full_name"], styles["Title"]))
    contact_line = " | ".join(filter(None, [resume.get("email"), resume.get("phone")]))
    if contact_line:
        elements.append(Paragraph(contact_line, styles["Normal"]))
    elements.append(Spacer(1, 12))

    if resume.get("summary"):
        elements.append(Paragraph("Summary", styles["Heading2"]))
        elements.append(Paragraph(resume["summary"], styles["Normal"]))
        elements.append(Spacer(1, 12))

    if resume.get("skills"):
        elements.append(Paragraph("Skills", styles["Heading2"]))
        elements.append(Paragraph(", ".join(resume["skills"]), styles["Normal"]))
        elements.append(Spacer(1, 12))

    if resume.get("experience"):
        elements.append(Paragraph("Experience", styles["Heading2"]))
        for job in resume["experience"]:
            title_line = f"{job.get('title', '')} — {job.get('company', '')}"
            elements.append(Paragraph(title_line, styles["Heading3"]))
            if job.get("duration"):
                elements.append(Paragraph(job["duration"], styles["Normal"]))
            for bullet in job.get("bullets", []):
                elements.append(Paragraph(f"• {bullet}", styles["Normal"]))
            elements.append(Spacer(1, 8))

    if resume.get("education"):
        elements.append(Paragraph("Education", styles["Heading2"]))
        for edu in resume["education"]:
            elements.append(Paragraph(edu, styles["Normal"]))

    doc.build(elements)
    buffer.seek(0)
    return buffer